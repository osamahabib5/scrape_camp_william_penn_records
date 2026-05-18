from __future__ import annotations

import argparse
import io
import json
import os
import re
import zipfile
from pathlib import Path
from typing import Any

import fitz
import pandas as pd
from PIL import Image
from pypdf import PdfReader
import pytesseract
from pytesseract import TesseractNotFoundError

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover - import should exist in this workspace
    raise RuntimeError("python-docx is required to read Word documents.") from exc


OUTPUT_FILE_NAME = "camp_william_penn_civil_war_troops.xlsx"
CHECKPOINT_FILE_NAME = "camp_william_penn_checkpoint.json"
SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]
FOLDER_MIME_TYPE = "application/vnd.google-apps.folder"
GOOGLE_DOC_MIME_TYPE = "application/vnd.google-apps.document"
PDF_MIME_TYPE = "application/pdf"
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOC_MIME_TYPE = "application/msword"
JPEG_MIME_TYPES = {"image/jpeg", "image/jpg"}
PNG_MIME_TYPE = "image/png"
TARGET_UNIT_FOLDERS = [
    "127th USCT",
    "22nd USCT",
    "24th USCT",
    "25th USCT",
    "32nd USCT",
    "3rd USCT",
    "41st USCT",
    "43rd USCT",
    "45th USCT",
    "6th USCT",
    "8th USCT",
]
DEFAULT_RESUME_UNITS = [
    "6th USCT",
    "127th USCT",
    "22nd USCT",
    "24th USCT",
    "25th USCT",
    "32nd USCT",
    "41st USCT",
    "43rd USCT",
    "45th USCT",
    "8th USCT",
]
ENLISTED_LABEL = "Enlisted_Men"
OFFICER_LABEL = "Officers"
PERSONNEL_FOLDER_MAP = {
    "ENLISTED MEN": ENLISTED_LABEL,
    "OFFICERS": OFFICER_LABEL,
}


def ensure_google_dependencies() -> None:
    try:
        import google.auth.transport.requests  # noqa: F401
        import google.oauth2.credentials  # noqa: F401
        import google_auth_oauthlib.flow  # noqa: F401
        import googleapiclient.discovery  # noqa: F401
        import googleapiclient.http  # noqa: F401
    except ImportError as exc:
        raise RuntimeError(
            "Google Drive packages are missing. Install them with: "
            "pip install google-api-python-client google-auth-httplib2 google-auth-oauthlib"
        ) from exc


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.split("\n")]
    collapsed: list[str] = []
    blank_streak = 0

    for line in lines:
        stripped = line.strip()
        if not stripped:
            blank_streak += 1
            if blank_streak <= 1:
                collapsed.append("")
            continue

        blank_streak = 0
        collapsed.append(re.sub(r"[ \t]+", " ", stripped))

    return "\n".join(collapsed).strip()


def sanitize_column_name(name: str) -> str:
    sanitized = re.sub(r"\s+", "_", name.strip())
    sanitized = re.sub(r"[^\w]", "_", sanitized)
    sanitized = re.sub(r"_+", "_", sanitized)
    return sanitized.strip("_")


def split_full_name(full_name: str) -> tuple[str, str]:
    cleaned_name = normalize_text(full_name).replace("\n", " ").strip()
    if not cleaned_name:
        return "", ""

    if "," in cleaned_name:
        surname, first_name = cleaned_name.split(",", 1)
        return first_name.strip(), surname.strip()

    parts = cleaned_name.split()
    if len(parts) == 1:
        return "", parts[0]

    return " ".join(parts[:-1]).strip(), parts[-1].strip()


def normalize_label(label: str) -> str:
    return re.sub(r"[^A-Z0-9]+", "", label.upper())


def natural_sort_key(value: str) -> list[Any]:
    key: list[tuple[int, Any]] = []
    for part in re.split(r"(\d+)", value):
        if not part:
            continue
        if part.isdigit():
            key.append((0, int(part)))
        else:
            key.append((1, part.lower()))
    return key


def shutil_which(command: str) -> str | None:
    path_value = os.environ.get("PATH", "")
    for folder in path_value.split(os.pathsep):
        if not folder:
            continue

        candidate = Path(folder) / f"{command}.exe"
        if candidate.exists():
            return str(candidate)

        plain_candidate = Path(folder) / command
        if plain_candidate.exists():
            return str(plain_candidate)

    return None


def find_tesseract_executable() -> str | None:
    candidates = [
        os.environ.get("TESSERACT_CMD"),
        r"D:\Tesseract\tesseract.exe",
        shutil_which("tesseract"),
    ]

    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return candidate

    return None


def configure_tesseract_from_path(executable_path: str | None) -> bool:
    if executable_path:
        candidate = Path(executable_path).expanduser().resolve()
        if not candidate.exists():
            raise FileNotFoundError(f"Tesseract executable not found: {candidate}")
        pytesseract.pytesseract.tesseract_cmd = str(candidate)
        return True

    executable = find_tesseract_executable()
    if not executable:
        return False

    pytesseract.pytesseract.tesseract_cmd = executable
    return True


def ocr_pdf_page(pdf_bytes: bytes, page_number: int, language: str) -> str:
    with fitz.open(stream=pdf_bytes, filetype="pdf") as pdf_document:
        page = pdf_document.load_page(page_number)
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        image_bytes = pixmap.tobytes("png")

    with Image.open(io.BytesIO(image_bytes)) as image:
        return normalize_text(pytesseract.image_to_string(image, lang=language))


def extract_source_information_from_image(image_bytes: bytes, language: str) -> dict[str, str]:
    with Image.open(io.BytesIO(image_bytes)) as image:
        text = normalize_text(pytesseract.image_to_string(image, lang=language))
    return parse_source_information(text)


def extract_pdf_pages(pdf_bytes: bytes, tesseract_cmd: str | None, language: str) -> list[str]:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    ocr_available = configure_tesseract_from_path(tesseract_cmd)
    extracted_pages: list[str] = []

    for page_number, page in enumerate(reader.pages):
        embedded_text = normalize_text(page.extract_text() or "")
        if embedded_text:
            extracted_pages.append(embedded_text)
            continue

        if not ocr_available:
            extracted_pages.append("")
            continue

        try:
            extracted_pages.append(ocr_pdf_page(pdf_bytes, page_number, language))
        except TesseractNotFoundError:
            extracted_pages.append("")
            ocr_available = False

    return extracted_pages


def extract_docx_text(docx_bytes: bytes) -> str:
    try:
        document = Document(io.BytesIO(docx_bytes))
    except (zipfile.BadZipFile, Exception) as exc:
        raise RuntimeError(f"Failed to parse DOCX file: {exc}") from exc
    
    text_blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if text:
            text_blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = normalize_text(cell.text)
                if text:
                    text_blocks.append(text)

    return "\n".join(text_blocks)


def parse_source_information(text: str) -> dict[str, str]:
    if "SOURCE INFORMATION" not in text:
        return {}

    lines = [line.strip() for line in text.splitlines()]
    fields: dict[str, str] = {}
    current_key: str | None = None
    current_value: list[str] = []

    for line in lines:
        if not line or line == "SOURCE INFORMATION":
            continue

        if line.endswith(":"):
            if current_key:
                fields[current_key] = " ".join(current_value).strip()
            current_key = line[:-1]
            current_value = []
            continue

        if current_key:
            current_value.append(line)

    if current_key:
        fields[current_key] = " ".join(current_value).strip()

    return fields


def merge_fields(base_fields: dict[str, str], new_fields: dict[str, str]) -> dict[str, str]:
    merged = dict(base_fields)
    for key, value in new_fields.items():
        if value and not merged.get(key):
            merged[key] = value
    return merged


def extract_source_information_from_pdf(pdf_bytes: bytes, tesseract_cmd: str | None, language: str) -> dict[str, str]:
    pages = extract_pdf_pages(pdf_bytes, tesseract_cmd, language)
    for text in pages:
        fields = parse_source_information(text)
        if fields:
            return fields
    return {}


def extract_source_information_from_docx(docx_bytes: bytes) -> dict[str, str]:
    text = extract_docx_text(docx_bytes)
    return parse_source_information(text)


def extract_file_extension(file_metadata: dict[str, Any]) -> str:
    file_name = file_metadata.get("name", "")
    if "." not in file_name:
        return ""
    return file_name.rsplit(".", 1)[-1].lower()


def escape_drive_query_value(value: str) -> str:
    return value.replace("\\", "\\\\").replace("'", "\\'")


def parse_drive_folder_id(folder_reference: str) -> str:
    reference = folder_reference.strip()
    match = re.search(r"/folders/([a-zA-Z0-9_-]+)", reference)
    if match:
        return match.group(1)
    return reference


def get_drive_service(credentials_path: Path, token_path: Path):
    ensure_google_dependencies()

    from google.auth.transport.requests import Request
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from googleapiclient.discovery import build

    creds = None
    if token_path.exists():
        creds = Credentials.from_authorized_user_file(str(token_path), SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not credentials_path.exists():
                raise FileNotFoundError(
                    f"Google OAuth credentials file not found: {credentials_path}"
                )
            flow = InstalledAppFlow.from_client_secrets_file(str(credentials_path), SCOPES)
            creds = flow.run_local_server(port=0)

        token_path.write_text(creds.to_json(), encoding="utf-8")

    return build("drive", "v3", credentials=creds)


def list_drive_children(service, parent_id: str, mime_type: str | None = None) -> list[dict[str, Any]]:
    query_parts = [f"'{parent_id}' in parents", "trashed = false"]
    if mime_type:
        query_parts.append(f"mimeType = '{mime_type}'")

    query = " and ".join(query_parts)
    fields = "nextPageToken, files(id, name, mimeType, fileExtension, modifiedTime)"
    page_token = None
    children: list[dict[str, Any]] = []

    while True:
        response = (
            service.files()
            .list(
                q=query,
                pageSize=1000,
                fields=fields,
                supportsAllDrives=True,
                includeItemsFromAllDrives=True,
                pageToken=page_token,
            )
            .execute()
        )
        children.extend(response.get("files", []))
        page_token = response.get("nextPageToken")
        if not page_token:
            break

    return sorted(children, key=lambda item: natural_sort_key(item["name"]))


def find_named_child_folder(service, parent_id: str, target_name: str) -> dict[str, Any] | None:
    escaped_name = escape_drive_query_value(target_name)
    query = (
        f"'{parent_id}' in parents and trashed = false and "
        f"mimeType = '{FOLDER_MIME_TYPE}' and name = '{escaped_name}'"
    )
    response = (
        service.files()
        .list(
            q=query,
            pageSize=10,
            fields="files(id, name, mimeType)",
            supportsAllDrives=True,
            includeItemsFromAllDrives=True,
        )
        .execute()
    )
    files = response.get("files", [])
    if files:
        return files[0]

    normalized_target = normalize_label(target_name)
    for folder in list_drive_children(service, parent_id, mime_type=FOLDER_MIME_TYPE):
        if normalize_label(folder["name"]) == normalized_target:
            return folder

    return None


def list_soldier_folders(service, personnel_folder_id: str, limit: int | None) -> list[dict[str, Any]]:
    folders = list_drive_children(service, personnel_folder_id, mime_type=FOLDER_MIME_TYPE)
    if limit is None:
        return folders
    return folders[:limit]


def download_blob_file(service, file_id: str) -> bytes:
    from googleapiclient.http import MediaIoBaseDownload

    request = service.files().get_media(fileId=file_id, supportsAllDrives=True)
    buffer = io.BytesIO()
    downloader = MediaIoBaseDownload(buffer, request)
    done = False

    while not done:
        _, done = downloader.next_chunk()

    return buffer.getvalue()


def export_google_doc_as_docx(service, file_id: str) -> bytes:
    from googleapiclient.http import MediaIoBaseDownload

    request = service.files().export_media(fileId=file_id, mimeType=DOCX_MIME_TYPE)
    buffer = io.BytesIO()
    downloader = MediaIoBaseDownload(buffer, request)
    done = False

    while not done:
        _, done = downloader.next_chunk()

    return buffer.getvalue()


def determine_supported_document_type(file_metadata: dict[str, Any]) -> str | None:
    mime_type = file_metadata.get("mimeType", "")
    extension = extract_file_extension(file_metadata)

    if mime_type == PDF_MIME_TYPE or extension == "pdf":
        return "pdf"
    if mime_type in {DOCX_MIME_TYPE, GOOGLE_DOC_MIME_TYPE} or extension == "docx":
        return "docx"
    if mime_type == DOC_MIME_TYPE or extension == "doc":
        return "doc"
    if mime_type in JPEG_MIME_TYPES or extension in {"jpg", "jpeg"}:
        return "image"
    if mime_type == PNG_MIME_TYPE or extension == "png":
        return "image"
    return None


def extract_source_information_from_drive_file(
    service,
    file_metadata: dict[str, Any],
    tesseract_cmd: str | None,
    language: str,
) -> dict[str, str]:
    file_name = file_metadata.get("name", "unknown")
    document_type = determine_supported_document_type(file_metadata)
    file_id = file_metadata["id"]

    try:
        if document_type == "pdf":
            file_bytes = download_blob_file(service, file_id)
            return extract_source_information_from_pdf(file_bytes, tesseract_cmd, language)

        if document_type == "docx":
            if file_metadata.get("mimeType") == GOOGLE_DOC_MIME_TYPE:
                file_bytes = export_google_doc_as_docx(service, file_id)
            else:
                file_bytes = download_blob_file(service, file_id)
            return extract_source_information_from_docx(file_bytes)

        if document_type == "doc":
            print(f"Skipping unsupported legacy Word file: {file_name}")
            return {}

        if document_type == "image":
            if not configure_tesseract_from_path(tesseract_cmd):
                print(f"Skipping image OCR because Tesseract is not configured: {file_name}")
                return {}
            file_bytes = download_blob_file(service, file_id)
            return extract_source_information_from_image(file_bytes, language)

        return {}
    except (zipfile.BadZipFile, Exception) as exc:
        print(f"Warning: Failed to extract text from file '{file_name}': {exc}")
        return {}


def build_record(
    unit_name: str,
    personnel_label: str,
    soldier_folder_name: str,
    source_files: list[str],
    fields: dict[str, str],
) -> dict[str, str]:
    record: dict[str, str] = {
        "Unit": unit_name,
        "Officer_Enlisted_Men": personnel_label,
        "Soldier_Folder_Name": soldier_folder_name,
        "Source_Files": "; ".join(source_files),
    }

    for key, value in fields.items():
        if key == "Full Name":
            first_name, surname = split_full_name(value)
            record["FirstName"] = first_name
            record["Surname"] = surname
            continue

        record[sanitize_column_name(key)] = value

    record.setdefault("FirstName", "")
    record.setdefault("Surname", "")
    return record


def collect_soldier_record(
    service,
    soldier_folder: dict[str, Any],
    unit_name: str,
    personnel_label: str,
    tesseract_cmd: str | None,
    language: str,
) -> dict[str, str]:
    fields: dict[str, str] = {}
    source_files: list[str] = []

    for file_metadata in list_drive_children(service, soldier_folder["id"]):
        if file_metadata.get("mimeType") == FOLDER_MIME_TYPE:
            continue

        document_type = determine_supported_document_type(file_metadata)
        if document_type is None:
            continue

        source_files.append(file_metadata["name"])
        extracted_fields = extract_source_information_from_drive_file(
            service,
            file_metadata,
            tesseract_cmd,
            language,
        )
        fields = merge_fields(fields, extracted_fields)

    return build_record(
        unit_name=unit_name,
        personnel_label=personnel_label,
        soldier_folder_name=soldier_folder["name"],
        source_files=source_files,
        fields=fields,
    )


def collect_drive_records(
    service,
    root_folder_id: str,
    units: list[str],
    personnel_targets: list[str],
    soldier_limit: int | None,
    tesseract_cmd: str | None,
    language: str,
    output_path: Path,
    checkpoint_path: Path,
) -> dict[str, dict[str, int]]:
    existing_dataframe = load_existing_dataframe(output_path)
    processed_keys = extract_processed_keys(existing_dataframe)
    next_id = get_next_record_id(existing_dataframe)
    summary: dict[str, dict[str, int]] = {}

    for unit_name in units:
        unit_folder = find_named_child_folder(service, root_folder_id, unit_name)
        if not unit_folder:
            print(f"Unit folder not found under root: {unit_name}")
            continue

        for personnel_name in personnel_targets:
            personnel_folder = find_named_child_folder(service, unit_folder["id"], personnel_name)
            if not personnel_folder:
                print(f"Personnel folder not found under {unit_name}: {personnel_name}")
                continue

            soldier_folders = list_soldier_folders(service, personnel_folder["id"], soldier_limit)
            for soldier_folder in soldier_folders:
                record_key = soldier_record_key(
                    unit_name,
                    PERSONNEL_FOLDER_MAP[personnel_name],
                    soldier_folder["name"],
                )
                if record_key in processed_keys:
                    print(f"Skipping already processed record: {unit_name} -> {personnel_name} -> {soldier_folder['name']}")
                    continue

                print(f"Processing {unit_name} -> {personnel_name} -> {soldier_folder['name']}")
                save_checkpoint(
                    checkpoint_path,
                    unit_name=unit_name,
                    personnel_name=personnel_name,
                    soldier_folder_name=soldier_folder["name"],
                    output_path=output_path,
                )
                record = collect_soldier_record(
                    service=service,
                    soldier_folder=soldier_folder,
                    unit_name=unit_name,
                    personnel_label=PERSONNEL_FOLDER_MAP[personnel_name],
                    tesseract_cmd=tesseract_cmd,
                    language=language,
                )
                next_id = append_record_to_excel(output_path, record, next_id)
                processed_keys.add(record_key)
                unit_summary = summary.setdefault(unit_name, {})
                personnel_label = PERSONNEL_FOLDER_MAP[personnel_name]
                unit_summary[personnel_label] = unit_summary.get(personnel_label, 0) + 1

    return summary


def build_dataframe(records: list[dict[str, str]]) -> pd.DataFrame:
    all_columns: list[str] = [
        "Id",
        "Unit",
        "Officer_Enlisted_Men",
        "Soldier_Folder_Name",
        "Source_Files",
        "FirstName",
        "Surname",
    ]
    seen_columns = set(all_columns)

    discovered_columns = sorted(
        {
            column
            for record in records
            for column in record.keys()
            if column not in seen_columns
        }
    )
    all_columns.extend(discovered_columns)

    normalized_records: list[dict[str, str | int]] = []
    for index, record in enumerate(records, start=1):
        row: dict[str, str | int] = {"Id": index}
        for column in all_columns:
            if column == "Id":
                continue
            row[column] = record.get(column, "")
        normalized_records.append(row)

    return pd.DataFrame(normalized_records, columns=all_columns)


def soldier_record_key(unit_name: str, personnel_label: str, soldier_folder_name: str) -> tuple[str, str, str]:
    return (unit_name, personnel_label, soldier_folder_name)


def load_existing_dataframe(output_path: Path) -> pd.DataFrame:
    if not output_path.exists():
        return pd.DataFrame()

    dataframe = pd.read_excel(output_path)
    for column in ["Unit", "Officer_Enlisted_Men", "Soldier_Folder_Name"]:
        if column not in dataframe.columns:
            dataframe[column] = ""
    return dataframe


def extract_processed_keys(dataframe: pd.DataFrame) -> set[tuple[str, str, str]]:
    if dataframe.empty:
        return set()

    processed_keys: set[tuple[str, str, str]] = set()
    rows = dataframe[["Unit", "Officer_Enlisted_Men", "Soldier_Folder_Name"]].fillna("")
    for row in rows.itertuples(index=False):
        processed_keys.add(soldier_record_key(row[0], row[1], row[2]))
    return processed_keys


def get_next_record_id(dataframe: pd.DataFrame) -> int:
    if dataframe.empty or "Id" not in dataframe.columns:
        return 1

    id_series = pd.to_numeric(dataframe["Id"], errors="coerce").dropna()
    if id_series.empty:
        return 1
    return int(id_series.max()) + 1


def merge_dataframes(existing_dataframe: pd.DataFrame, new_records: list[dict[str, str | int]]) -> pd.DataFrame:
    new_dataframe = pd.DataFrame(new_records)
    if existing_dataframe.empty:
        combined = new_dataframe
    else:
        combined = pd.concat([existing_dataframe, new_dataframe], ignore_index=True)

    preferred_columns = [
        "Id",
        "Unit",
        "Officer_Enlisted_Men",
        "Soldier_Folder_Name",
        "Source_Files",
        "FirstName",
        "Surname",
    ]
    remaining_columns = [column for column in combined.columns if column not in preferred_columns]
    ordered_columns = preferred_columns + sorted(remaining_columns)
    return combined.reindex(columns=ordered_columns)


def append_record_to_excel(output_path: Path, record: dict[str, str], next_id: int) -> int:
    existing_dataframe = load_existing_dataframe(output_path)
    row = {"Id": next_id, **record}
    combined_dataframe = merge_dataframes(existing_dataframe, [row])
    combined_dataframe.to_excel(output_path, index=False)
    return next_id + 1


def load_checkpoint(checkpoint_path: Path) -> dict[str, Any]:
    if not checkpoint_path.exists():
        return {}
    return json.loads(checkpoint_path.read_text(encoding="utf-8"))


def save_checkpoint(
    checkpoint_path: Path,
    *,
    unit_name: str,
    personnel_name: str,
    soldier_folder_name: str,
    output_path: Path,
) -> None:
    payload = {
        "unit_name": unit_name,
        "personnel_name": personnel_name,
        "soldier_folder_name": soldier_folder_name,
        "output_path": str(output_path),
        "status": "in_progress",
    }
    checkpoint_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def mark_checkpoint_complete(checkpoint_path: Path, output_path: Path) -> None:
    payload = {
        "status": "completed",
        "output_path": str(output_path),
    }
    checkpoint_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def print_processing_summary(summary: dict[str, dict[str, int]]) -> None:
    print("\nProcessing summary:")
    if not summary:
        print("  No new records were added.")
        return

    total_records = 0
    for unit_name in DEFAULT_RESUME_UNITS + ["3rd USCT"]:
        if unit_name not in summary:
            continue
        unit_summary = summary[unit_name]
        unit_total = sum(unit_summary.values())
        total_records += unit_total
        enlisted_count = unit_summary.get(ENLISTED_LABEL, 0)
        officer_count = unit_summary.get(OFFICER_LABEL, 0)
        print(
            f"  {unit_name}: {unit_total} added "
            f"({ENLISTED_LABEL}={enlisted_count}, {OFFICER_LABEL}={officer_count})"
        )

    remaining_units = [unit for unit in summary if unit not in DEFAULT_RESUME_UNITS + ["3rd USCT"]]
    for unit_name in sorted(remaining_units):
        unit_summary = summary[unit_name]
        unit_total = sum(unit_summary.values())
        total_records += unit_total
        enlisted_count = unit_summary.get(ENLISTED_LABEL, 0)
        officer_count = unit_summary.get(OFFICER_LABEL, 0)
        print(
            f"  {unit_name}: {unit_total} added "
            f"({ENLISTED_LABEL}={enlisted_count}, {OFFICER_LABEL}={officer_count})"
        )

    print(f"  Total new records added: {total_records}")


def resolve_output_directory(requested_directory: str | None) -> Path:
    if not requested_directory:
        return Path.cwd().resolve()

    directory = Path(requested_directory).expanduser().resolve()
    if not directory.exists():
        raise FileNotFoundError(f"Output directory not found: {directory}")
    if not directory.is_dir():
        raise NotADirectoryError(f"Output path is not a directory: {directory}")
    return directory


def parse_personnel_targets(selection: str) -> list[str]:
    normalized = selection.strip().upper()
    if normalized == "BOTH":
        return list(PERSONNEL_FOLDER_MAP.keys())
    if normalized not in PERSONNEL_FOLDER_MAP:
        raise ValueError(
            "personnel type must be one of: ENLISTED MEN, OFFICERS, BOTH"
        )
    return [normalized]


def create_excel_from_drive(
    root_folder_reference: str,
    output_directory: Path,
    credentials_path: Path,
    token_path: Path,
    units: list[str],
    personnel_targets: list[str],
    soldier_limit: int | None,
    tesseract_cmd: str | None,
    language: str,
    checkpoint_path: Path,
) -> Path:
    service = get_drive_service(credentials_path, token_path)
    root_folder_id = parse_drive_folder_id(root_folder_reference)
    output_path = output_directory / OUTPUT_FILE_NAME
    summary = collect_drive_records(
        service=service,
        root_folder_id=root_folder_id,
        units=units,
        personnel_targets=personnel_targets,
        soldier_limit=soldier_limit,
        tesseract_cmd=tesseract_cmd,
        language=language,
        output_path=output_path,
        checkpoint_path=checkpoint_path,
    )
    if load_existing_dataframe(output_path).empty:
        raise RuntimeError("No soldier records were collected from Google Drive.")

    mark_checkpoint_complete(checkpoint_path, output_path)
    print_processing_summary(summary)
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(
        description=(
            "Authenticate to Google Drive, crawl the Camp William Penn folder hierarchy, "
            "extract SOURCE INFORMATION from PDFs and Word documents, and save the "
            "consolidated results into one Excel file."
        )
    )
    parser.add_argument(
        "--drive-root",
        required=True,
        help="Google Drive root folder ID or full folder URL.",
    )
    parser.add_argument(
        "--credentials-file",
        default="credentials.json",
        help="Path to the Google OAuth desktop-app credentials JSON file.",
    )
    parser.add_argument(
        "--token-file",
        default="token.json",
        help="Path where the Google OAuth access token should be stored.",
    )
    parser.add_argument(
        "--output-dir",
        help="Directory where the Excel file should be written. Defaults to the current directory.",
    )
    parser.add_argument(
        "--checkpoint-file",
        default=CHECKPOINT_FILE_NAME,
        help="Path to the checkpoint JSON file. Default: camp_william_penn_checkpoint.json",
    )
    parser.add_argument(
        "--personnel-type",
        default="BOTH",
        help="ENLISTED MEN, OFFICERS, or BOTH. Default: BOTH",
    )
    parser.add_argument(
        "--unit",
        action="append",
        dest="units",
        help=(
            "Unit folder to process. Repeat the flag to process more than one unit. "
            "If omitted, the script resumes from 6th USCT and then processes the remaining configured units."
        ),
    )
    parser.add_argument(
        "--soldier-limit",
        type=int,
        default=0,
        help=(
            "Maximum number of soldier folders to process per personnel folder. "
            "Default: 0, which processes all soldier folders."
        ),
    )
    parser.add_argument(
        "--all-units",
        action="store_true",
        help="Process all configured unit folders instead of only the default test unit.",
    )
    parser.add_argument(
        "--tesseract-cmd",
        help="Optional full path to tesseract.exe if it is not on PATH.",
    )
    parser.add_argument(
        "--ocr-lang",
        default="eng",
        help="OCR language code for Tesseract, such as 'eng'. Default: eng",
    )
    args = parser.parse_args()

    output_directory = resolve_output_directory(args.output_dir)
    credentials_path = Path(args.credentials_file).expanduser().resolve()
    token_path = Path(args.token_file).expanduser().resolve()
    checkpoint_path = Path(args.checkpoint_file).expanduser().resolve()

    if args.units:
        units = args.units
    elif args.all_units:
        units = TARGET_UNIT_FOLDERS
    else:
        units = DEFAULT_RESUME_UNITS

    soldier_limit = None if args.soldier_limit == 0 else args.soldier_limit
    personnel_targets = parse_personnel_targets(args.personnel_type)

    output_path = create_excel_from_drive(
        root_folder_reference=args.drive_root,
        output_directory=output_directory,
        credentials_path=credentials_path,
        token_path=token_path,
        units=units,
        personnel_targets=personnel_targets,
        soldier_limit=soldier_limit,
        tesseract_cmd=args.tesseract_cmd,
        language=args.ocr_lang,
        checkpoint_path=checkpoint_path,
    )
    print(f"Created Excel file: {output_path}")


if __name__ == "__main__":
    main()
