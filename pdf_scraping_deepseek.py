"""
pdf_ocr_to_docx.py
------------------
Extracts handwritten/printed content from a scanned PDF using DeepSeek's
vision model (OCR), then saves the extracted text into a Word (.docx) file.

Requirements:
    pip install openai python-docx pymupdf

Environment variables:
    DEEPSEEK_API_KEY  — your DeepSeek API key

NOTE ON DEEPSEEK VISION API:
  DeepSeek does NOT support image_url with base64 data URIs like OpenAI does.
  Images must be sent using the "image_url" type but with an HTTPS URL,
  OR using the correct base64 format:
    { "type": "image_url", "image_url": { "url": "data:image/jpeg;base64,..." } }
  If that also fails, the workaround is to upload images to a temporary host
  first, or use a different vision provider (e.g. Claude, GPT-4o) for OCR
  and just use DeepSeek for text tasks.
"""

import os
import sys
import base64
from pathlib import Path
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv()

try:
    import fitz  # PyMuPDF
except ImportError:
    sys.exit("Error: PyMuPDF not installed. Run: pip install pymupdf")

# ---------------------------------------------------------------------------
# CONFIGURE HERE
# ---------------------------------------------------------------------------

PDF_PATH = "Abercrombie_William_127th_USCT.pdf"            # <-- set your PDF path here
OUTPUT_DOCX = "extracted_forms.docx"  # <-- set your desired output path

DEEPSEEK_MODEL = "deepseek-v4-flash"
DEEPSEEK_BASE_URL = "https://api.deepseek.com"

# Resolution for rendering PDF pages (higher = more readable for OCR)
DPI = 200

OCR_SYSTEM_PROMPT = (
    "You are an expert OCR assistant specialising in handwritten and printed forms. "
    "When given an image of a form page, extract ALL visible text exactly as written, "
    "preserving the logical structure of the form. "
    "For each field, output it as:  Field Label: Value\n"
    "If the field is blank, write:  Field Label: [blank]\n"
    "Preserve section headings. Do NOT add commentary or extra explanation — "
    "output only the extracted content."
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def get_client() -> OpenAI:
    api_key = os.environ.get("DEEPSEEK_API_KEY")
    if not api_key:
        sys.exit("Error: DEEPSEEK_API_KEY environment variable is not set.")
    return OpenAI(api_key=api_key, base_url=DEEPSEEK_BASE_URL)


def extract_pdf_pages_as_images(pdf_path: str, dpi: int = DPI) -> list[tuple[int, bytes]]:
    """
    Render each PDF page to a JPEG using PyMuPDF.
    Works on Windows, macOS, and Linux — no external tools needed.
    Returns list of (page_number, jpeg_bytes).
    """
    doc = fitz.open(pdf_path)
    zoom = dpi / 72  # 72 DPI is PyMuPDF's base resolution
    pages = []
    for i, page in enumerate(doc, start=1):
        pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
        pages.append((i, pix.tobytes(output="jpeg")))
    doc.close()
    return pages


def ocr_page(client: OpenAI, jpeg_bytes: bytes, page_num: int) -> str:
    """
    Send one page image to DeepSeek vision and return extracted text.

    DeepSeek uses the OpenAI-compatible API but requires the base64 image
    to be embedded in a data URI within the image_url field.
    """
    print(f"  OCR-ing page {page_num} ...", flush=True)

    b64 = base64.standard_b64encode(jpeg_bytes).decode("utf-8")
    data_uri = f"data:image/jpeg;base64,{b64}"

    response = client.chat.completions.create(
        model=DEEPSEEK_MODEL,
        messages=[
            {"role": "system", "content": OCR_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": [
                    # DeepSeek vision format: image_url with data URI
                    {
                        "type": "image_url",
                        "image_url": {"url": data_uri},
                    },
                    {
                        "type": "text",
                        "text": (
                            f"This is page {page_num} of the scanned form. "
                            "Please extract all text and form field values as instructed."
                        ),
                    },
                ],
            },
        ],
        stream=False,
        max_tokens=4096,
    )

    return response.choices[0].message.content or ""


# ---------------------------------------------------------------------------
# Word document builder
# ---------------------------------------------------------------------------

def build_docx(pages_text: list[tuple[int, str]], output_path: str) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_heading("Extracted Form Content", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for page_num, text in pages_text:
        doc.add_heading(f"Page {page_num}", level=1)

        if text.strip():
            for line in text.splitlines():
                stripped = line.strip()
                if not stripped:
                    doc.add_paragraph()
                    continue
                para = doc.add_paragraph()
                if ":" in stripped:
                    colon_idx = stripped.index(":")
                    para.add_run(stripped[: colon_idx + 1]).bold = True
                    para.add_run(stripped[colon_idx + 1 :])
                else:
                    para.add_run(stripped)
        else:
            doc.add_paragraph("[No text extracted from this page]")

        if page_num < pages_text[-1][0]:
            doc.add_page_break()

    doc.save(output_path)
    print(f"\nWord document saved -> {output_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if not os.path.isfile(PDF_PATH):
        sys.exit(f"Error: PDF not found: {PDF_PATH}")

    print(f"Input  PDF : {PDF_PATH}")
    print(f"Output DOCX: {OUTPUT_DOCX}")

    client = get_client()

    print(f"\nRendering PDF pages at {DPI} DPI ...")
    pages = extract_pdf_pages_as_images(PDF_PATH, DPI)
    print(f"Found {len(pages)} page(s).\n")

    pages_text: list[tuple[int, str]] = []
    for page_num, jpeg_bytes in pages:
        try:
            text = ocr_page(client, jpeg_bytes, page_num)
        except Exception as exc:
            print(f"  Warning: OCR failed for page {page_num}: {exc}")
            text = f"[OCR error on page {page_num}: {exc}]"
        pages_text.append((page_num, text))

    print("\nBuilding Word document ...")
    build_docx(pages_text, OUTPUT_DOCX)


if __name__ == "__main__":
    main()