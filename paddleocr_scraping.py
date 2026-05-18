from __future__ import annotations
import os

# PERFORMANCE & STABILITY FLAGS
# Disabling MKLDNN prevents the PIR executor error on Windows CPU
os.environ["FLAGS_use_mkldnn"] = "0"
os.environ["PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK"] = "True"
os.environ["PADDLE_PDX_ENABLE_MKLDNN_BYDEFAULT"] = "0" 

import argparse
import contextlib
from io import BytesIO
import json
import re
from pathlib import Path
import subprocess
import sys
import tempfile
from collections.abc import Iterable

import fitz  # PyMuPDF
import numpy as np
from docx import Document
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
from pypdf import PdfReader
import pytesseract
from pytesseract import TesseractNotFoundError


TEST_PAGE_NUMBERS = [3]
OCR_TIMEOUT_SECONDS = 180


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.split("\n")]
    collapsed: list[str] = []
    blank_streak = 0

    for line in lines:
        stripped = line.strip()
        if not stripped:
            blank_streak += 1
            if blank_streak <= 1:
                collapsed.append("")
            continue

        blank_streak = 0
        collapsed.append(re.sub(r"[ \t]+", " ", stripped))

    return "\n".join(collapsed).strip()


def shutil_which(command: str) -> str | None:
    path_value = os.environ.get("PATH", "")
    for folder in path_value.split(os.pathsep):
        if not folder:
            continue

        candidate = Path(folder) / f"{command}.exe"
        if candidate.exists():
            return str(candidate)

        plain_candidate = Path(folder) / command
        if plain_candidate.exists():
            return str(plain_candidate)

    return None


def find_tesseract_executable() -> str | None:
    candidates = [
        os.environ.get("TESSERACT_CMD"),
        r"D:\Tesseract\tesseract.exe",
        shutil_which("tesseract"),
    ]

    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return candidate

    return None


def configure_tesseract() -> bool:
    executable = find_tesseract_executable()
    if not executable:
        return False

    pytesseract.pytesseract.tesseract_cmd = executable
    return True


def get_paddleocr():
    try:
        from paddleocr import PaddleOCR
    except ImportError as exc:
        raise RuntimeError(
            "PaddleOCR is not installed. Install it with: pip install paddleocr"
        ) from exc

    return PaddleOCR(
        use_doc_orientation_classify=False,
        use_doc_unwarping=False,
        use_textline_orientation=False,
        text_detection_model_name="PP-OCRv5_mobile_det",
        text_recognition_model_name="en_PP-OCRv5_mobile_rec",
        lang="en",
    )


def render_pdf_page(pdf_path: Path, page_number: int) -> Image.Image:
    """
    Renders a PDF page into a PIL Image.
    For archival forms and handwriting, we prefer higher resolution over speed.
    """
    with fitz.open(pdf_path) as pdf_document:
        page = pdf_document.load_page(page_number)
        pixmap = page.get_pixmap(matrix=fitz.Matrix(4, 4), alpha=False)
        png_bytes = pixmap.tobytes("png")

    with Image.open(BytesIO(png_bytes)) as image:
        return image.convert("RGB")


def extract_embedded_text(pdf_path: Path) -> list[str]:
    reader = PdfReader(str(pdf_path))
    return [normalize_text(page.extract_text() or "") for page in reader.pages]


def crop_document_region(image: Image.Image) -> Image.Image:
    grayscale = ImageOps.grayscale(image)
    pixel_array = np.array(grayscale)
    bright_pixels = np.argwhere(pixel_array > 150)
    if bright_pixels.size == 0:
        return image

    top, left = bright_pixels.min(axis=0)
    bottom, right = bright_pixels.max(axis=0)
    padding = 20
    left = max(0, left - padding)
    top = max(0, top - padding)
    right = min(image.width, right + padding)
    bottom = min(image.height, bottom + padding)

    cropped = image.crop((left, top, right, bottom))
    return ImageOps.expand(cropped, border=16, fill="white")


def build_page_regions(image: Image.Image) -> list[tuple[str, Image.Image]]:
    cropped = crop_document_region(image)
    height = cropped.height

    regions = [("full_page", cropped)]
    sections = [
        ("top_section", 0.00, 0.45),
        ("middle_section", 0.28, 0.74),
        ("bottom_section", 0.56, 1.00),
    ]
    for name, start_ratio, end_ratio in sections:
        top = int(height * start_ratio)
        bottom = int(height * end_ratio)
        section = cropped.crop((0, top, cropped.width, max(top + 1, bottom)))
        section = ImageOps.expand(section, border=20, fill="white")
        regions.append((name, section))

    focused_regions = [
        ("header_block", 0.00, 0.24),
        ("name_company_block", 0.06, 0.34),
        ("description_block", 0.22, 0.58),
        ("enlistment_block", 0.50, 0.78),
        ("signature_block", 0.82, 1.00),
    ]
    for name, start_ratio, end_ratio in focused_regions:
        top = int(height * start_ratio)
        bottom = int(height * end_ratio)
        region = cropped.crop((0, top, cropped.width, max(top + 1, bottom)))
        region = ImageOps.expand(region, border=20, fill="white")
        regions.append((name, region))

    return regions


def prepare_ocr_variants(image: Image.Image) -> list[tuple[str, Image.Image]]:
    grayscale = ImageOps.grayscale(image)
    autocontrast = ImageOps.autocontrast(grayscale)
    sharpened = autocontrast.filter(ImageFilter.SHARPEN)
    high_contrast = ImageEnhance.Contrast(sharpened).enhance(2.2)

    thresholded = high_contrast.point(lambda pixel: 255 if pixel > 175 else 0)
    inverted_thresholded = ImageOps.invert(thresholded)
    enlarged = high_contrast.resize(
        (high_contrast.width * 2, high_contrast.height * 2),
        Image.Resampling.LANCZOS,
    )
    enlarged_thresholded = thresholded.resize(
        (thresholded.width * 2, thresholded.height * 2),
        Image.Resampling.LANCZOS,
    )

    return [
        ("rgb", image.convert("RGB")),
        ("high_contrast_gray", high_contrast.convert("RGB")),
        ("thresholded", thresholded.convert("RGB")),
        ("enlarged_gray", enlarged.convert("RGB")),
        ("enlarged_thresholded", enlarged_thresholded.convert("RGB")),
    ]


def score_ocr_text(text: str) -> tuple[int, int]:
    normalized = normalize_text(text)
    lines = [line for line in normalized.splitlines() if line.strip()]
    keyword_hits = 0
    for keyword in [
        "appears on",
        "company descriptive book",
        "description",
        "enlistment",
        "age",
        "height",
        "eyes",
        "hair",
        "where born",
        "occupation",
        "remarks",
        "source information",
    ]:
        if keyword in normalized.lower():
            keyword_hits += 1

    return (keyword_hits, len(lines))


def collect_text_fragments(value: object) -> list[str]:
    fragments: list[str] = []

    if value is None:
        return fragments

    if isinstance(value, str):
        normalized = normalize_text(value)
        if normalized:
            fragments.append(normalized)
        return fragments

    if isinstance(value, dict):
        direct_keys = ("text", "rec_text", "transcription")
        for key in direct_keys:
            direct_value = value.get(key)
            if isinstance(direct_value, str):
                normalized = normalize_text(direct_value)
                if normalized:
                    fragments.append(normalized)

        if fragments:
            return fragments

        nested_keys = (
            "items",
            "rec_texts",
            "prunedResult",
            "pruned_result",
            "res",
            "result",
            "ocrResults",
            "data",
        )
        for key in nested_keys:
            if key in value:
                fragments.extend(collect_text_fragments(value[key]))
        if fragments:
            return fragments

        for nested_value in value.values():
            fragments.extend(collect_text_fragments(nested_value))
        return fragments

    if isinstance(value, Iterable) and not isinstance(value, (bytes, bytearray)):
        for item in value:
            fragments.extend(collect_text_fragments(item))
        return fragments

    for attr_name in ("items", "rec_texts", "prunedResult", "pruned_result", "res", "result"):
        attr_value = getattr(value, attr_name, None)
        if attr_value is not None:
            fragments.extend(collect_text_fragments(attr_value))
            if fragments:
                return fragments

    return fragments


def extract_text_from_prediction_result(result: object) -> str:
    lines = collect_text_fragments(result)
    return "\n".join(lines).strip()


def ocr_with_tesseract(image: Image.Image) -> tuple[str, list[str]]:
    diagnostics: list[str] = []
    if not configure_tesseract():
        return "", ["tesseract: executable not configured"]

    collected_blocks: list[str] = []
    tesseract_config = "--psm 6"
    for region_name, region_image in build_page_regions(image):
        for variant_name, variant_image in prepare_ocr_variants(region_image):
            try:
                text = normalize_text(
                    pytesseract.image_to_string(variant_image, lang="eng", config=tesseract_config)
                )
            except TesseractNotFoundError:
                return "", ["tesseract: executable not found at runtime"]

            diagnostics.append(
                f"tesseract {region_name}__{variant_name}: "
                f"{len([line for line in text.splitlines() if line.strip()])} lines"
            )
            if text:
                collected_blocks.append(text)

    merged_lines: list[str] = []
    seen_lines: set[str] = set()
    prioritized_blocks = sorted(collected_blocks, key=score_ocr_text, reverse=True)
    for block in prioritized_blocks:
        for line in block.splitlines():
            normalized_line = normalize_text(line)
            if not normalized_line:
                continue
            dedupe_key = normalized_line.lower()
            if dedupe_key in seen_lines:
                continue
            seen_lines.add(dedupe_key)
            merged_lines.append(normalized_line)

    return "\n".join(merged_lines).strip(), diagnostics


def run_paddle_prediction(ocr_engine, image_path: Path) -> str:
    results = ocr_engine.predict(str(image_path))
    candidate_parts: list[str] = []
    for result in results or []:
        text = extract_text_from_prediction_result(result)
        if text:
            candidate_parts.append(text)

    return "\n".join(candidate_parts).strip()


def run_worker_from_variant_dir(variant_dir: Path) -> list[dict[str, str]]:
    ocr_engine = get_paddleocr()
    payload: list[dict[str, str]] = []

    for image_path in sorted(variant_dir.glob("*.png")):
        variant_name = image_path.stem
        text = run_paddle_prediction(ocr_engine, image_path)
        payload.append({"variant_name": variant_name, "text": text})

    return payload


def build_ocr_tasks(image: Image.Image) -> list[tuple[str, Image.Image]]:
    tasks: list[tuple[str, Image.Image]] = []
    region_priority = {
        "full_page",
        "header_block",
        "name_company_block",
        "description_block",
        "enlistment_block",
    }
    variant_priority = {
        "rgb",
        "high_contrast_gray",
        "enlarged_gray",
    }
    for region_name, region_image in build_page_regions(image):
        if region_name not in region_priority:
            continue
        for variant_name, variant_image in prepare_ocr_variants(region_image):
            if variant_name not in variant_priority:
                continue
            tasks.append((f"{region_name}__{variant_name}", variant_image))
    return tasks


def save_image_as_png(image: Image.Image, image_path: Path) -> None:
    image.save(image_path, format="PNG")


def run_ocr_variants_in_subprocess(variants: list[tuple[str, Image.Image]]) -> tuple[bool, object]:
    with tempfile.TemporaryDirectory(prefix="paddleocr_variants_", dir=str(Path.cwd())) as tmp_dir:
        variant_dir = Path(tmp_dir)
        for variant_name, variant_image in variants:
            image_path = variant_dir / f"{variant_name}.png"
            save_image_as_png(variant_image, image_path)

        command = [
            sys.executable,
            str(Path(__file__).resolve()),
            "--worker-variant-dir",
            str(variant_dir),
        ]
        try:
            completed = subprocess.run(
                command,
                capture_output=True,
                text=True,
                timeout=OCR_TIMEOUT_SECONDS,
            )
        except subprocess.TimeoutExpired:
            return False, f"worker timed out after {OCR_TIMEOUT_SECONDS} seconds"

        if completed.returncode != 0:
            stderr_text = normalize_text(completed.stderr or "")
            stdout_text = normalize_text(completed.stdout or "")
            message = stderr_text or stdout_text or f"worker crashed with exit code {completed.returncode}"
            return False, message

        try:
            return True, json.loads(completed.stdout)
        except json.JSONDecodeError:
            return False, "worker returned invalid OCR payload"


def ocr_image(image: Image.Image) -> tuple[str, list[str], str]:
    diagnostics: list[str] = []
    collected_blocks: list[str] = []

    variants = build_ocr_tasks(image)
    success, payload = run_ocr_variants_in_subprocess(variants)
    if not success:
        diagnostics.append(f"worker: {payload}")
        fallback_text, fallback_diagnostics = ocr_with_tesseract(image)
        diagnostics.extend(fallback_diagnostics)
        return fallback_text, diagnostics, (
            "Tesseract fallback" if fallback_text else "PaddleOCR attempted, no text found"
        )

    for item in payload:
        variant_name = item.get("variant_name", "unknown")
        candidate_text = item.get("text", "")
        diagnostics.append(
            f"{variant_name}: {len([line for line in candidate_text.splitlines() if line.strip()])} lines"
        )
        if candidate_text:
            collected_blocks.append(candidate_text)

    merged_lines: list[str] = []
    seen_lines: set[str] = set()
    prioritized_blocks = sorted(collected_blocks, key=score_ocr_text, reverse=True)
    for block in prioritized_blocks:
        for line in block.splitlines():
            normalized_line = normalize_text(line)
            if not normalized_line:
                continue
            dedupe_key = normalized_line.lower()
            if dedupe_key in seen_lines:
                continue
            seen_lines.add(dedupe_key)
            merged_lines.append(normalized_line)

    merged_text = "\n".join(merged_lines).strip()
    if merged_text:
        return merged_text, diagnostics, "PaddleOCR"

    fallback_text, fallback_diagnostics = ocr_with_tesseract(image)
    diagnostics.extend(fallback_diagnostics)
    return fallback_text, diagnostics, (
        "Tesseract fallback" if fallback_text else "PaddleOCR attempted, no text found"
    )


def extract_pdf_text_with_paddleocr(pdf_path: Path) -> tuple[list[dict[str, str]], list[str]]:
    """
    Extracts text from PDF. 
    Optimization: Limited to specific test pages for rapid testing.
    """
    all_embedded_pages = extract_embedded_text(pdf_path)

    selected_page_indices = [
        page_number - 1
        for page_number in TEST_PAGE_NUMBERS
        if 1 <= page_number <= len(all_embedded_pages)
    ]
    extracted_pages: list[dict[str, str]] = []
    methods: list[str] = []

    print(
        f"Processing test pages {', '.join(str(page_number) for page_number in TEST_PAGE_NUMBERS)} "
        f"of {pdf_path.name}..."
    )

    for page_number in selected_page_indices:
        embedded_text = all_embedded_pages[page_number]
        if embedded_text:
            extracted_pages.append({"page_number": str(page_number + 1), "text": embedded_text})
            methods.append("embedded text")
            print(f"  - Page {page_number + 1}: Found embedded text.")
            continue

        print(f"  - Page {page_number + 1}: Running PaddleOCR in isolated worker...")
        image = render_pdf_page(pdf_path, page_number)
        ocr_text, diagnostics, method_label = ocr_image(image)
        for diagnostic in diagnostics:
            print(f"    * {diagnostic}")
        extracted_pages.append({"page_number": str(page_number + 1), "text": ocr_text})
        methods.append(method_label)

    return extracted_pages, methods


def parse_source_information(text: str) -> dict[str, str]:
    if "SOURCE INFORMATION" not in text:
        return {}

    lines = [line.strip() for line in text.splitlines()]
    fields: dict[str, str] = {}
    current_key: str | None = None
    current_value: list[str] = []

    for line in lines:
        if not line or line == "SOURCE INFORMATION":
            continue

        if line.endswith(":"):
            if current_key:
                fields[current_key] = " ".join(current_value).strip()
            current_key = line[:-1]
            current_value = []
            continue

        if current_key:
            current_value.append(line)

    if current_key:
        fields[current_key] = " ".join(current_value).strip()

    return fields


def add_source_information(document: Document, fields: dict[str, str]) -> None:
    if not fields:
        return

    document.add_heading("Source Information", level=1)
    for key, value in fields.items():
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{key}: ").bold = True
        paragraph.add_run(value or "[No value extracted]")


def add_page_text(document: Document, pages: list[dict[str, str]], methods: list[str]) -> None:
    document.add_heading("Extracted Text By Page (Test Batch)", level=1)

    for page, method in zip(pages, methods, strict=True):
        document.add_heading(f"Page {page['page_number']}", level=2)
        document.add_paragraph(f"Extraction method: {method}")
        if page["text"]:
            document.add_paragraph(page["text"])
        else:
            document.add_paragraph("No text was extracted from this page.")


def find_source_information(pages: list[dict[str, str]]) -> dict[str, str]:
    for page in pages:
        fields = parse_source_information(page["text"])
        if fields:
            return fields
    return {}


def create_word_document(pdf_path: Path, pages: list[dict[str, str]], methods: list[str]) -> Path:
    output_path = pdf_path.with_name(f"{pdf_path.stem}_TEST_EXTRACT.docx")
    document = Document()
    document.add_heading(f"TEST EXTRACTION: {pdf_path.name}", level=0)

    document.add_paragraph(
        f"Source PDF: {pdf_path.name}\n"
        f"Test Scope: Pages {', '.join(str(page_number) for page_number in TEST_PAGE_NUMBERS)} processed.\n"
        f"Extraction method used: PaddleOCR with archival-image preprocessing"
    )

    source_fields = find_source_information(pages)
    add_source_information(document, source_fields)
    add_page_text(document, pages, methods)
    document.save(str(output_path))
    return output_path


def process_pdf(pdf_path: Path) -> Path:
    pages, methods = extract_pdf_text_with_paddleocr(pdf_path)
    return create_word_document(pdf_path, pages, methods)


def resolve_pdf_files(directory: Path, requested_pdf: str | None) -> list[Path]:
    if requested_pdf:
        pdf_path = Path(requested_pdf).expanduser().resolve()
        if not pdf_path.exists():
            raise FileNotFoundError(f"PDF not found: {pdf_path}")
        return [pdf_path]

    pdf_files = sorted(directory.glob("*.pdf"))
    if not pdf_files:
        raise FileNotFoundError(f"No PDF files were found in {directory}")
    return pdf_files


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Speed-Optimized PaddleOCR Test Scraper"
    )
    parser.add_argument("pdf_path", nargs="?", help="Optional path to a single PDF.")
    parser.add_argument(
        "--worker-variant-dir",
        help=argparse.SUPPRESS,
    )
    args = parser.parse_args()

    if args.worker_variant_dir:
        variant_dir = Path(args.worker_variant_dir).expanduser().resolve()
        with open(os.devnull, "w", encoding="utf-8") as devnull:
            with contextlib.redirect_stdout(devnull), contextlib.redirect_stderr(devnull):
                payload = run_worker_from_variant_dir(variant_dir)
        print(json.dumps(payload))
        return

    directory = Path.cwd().resolve()
    pdf_files = resolve_pdf_files(directory, args.pdf_path)
    for pdf_path in pdf_files:
        output_path = process_pdf(pdf_path)
        print(f"Successfully Created Test Word Doc: {output_path}")


if __name__ == "__main__":
    main()
